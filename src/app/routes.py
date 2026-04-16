# src/app/routes.py

import base64
import csv
import io
import json
import re
import sqlite3
import os
import statistics
from collections import defaultdict
from datetime import datetime, timezone
from flask import Blueprint, render_template, request, jsonify, redirect, url_for, session, Response
from werkzeug.security import generate_password_hash, check_password_hash

from . import study_schedule as sched

bp = Blueprint('main', __name__)


@bp.before_request
def require_session_unlock():
    """Require a signed-in user and PIN unlock for this browser session."""
    ep = request.endpoint
    if ep is None or ep == 'static':
        return

    if ep in (
        'main.home',
        'main.unlock_page',
        'main.api_onboarding_complete',
        'main.api_unlock',
        'main.api_logout',
        'main.api_dev_session',
    ):
        return

    if ep == 'main.api_results_upload':
        return

    uid = session.get('user_id')
    if not uid:
        if request.path.startswith('/api/'):
            return jsonify({'message': 'Not signed in.', 'code': 'auth_required'}), 401
        return redirect(url_for('main.home'))

    if session.get('pin_unlocked'):
        return

    if request.path.startswith('/api/'):
        return jsonify({'message': 'Enter your PIN to continue.', 'code': 'unlock_required'}), 401
    return redirect(url_for('main.unlock_page'))


def _db_path():
    """SQLite path; override with PERFORMANCELENS_DB_PATH for Tauri/Electron bundles."""
    return os.environ.get(
        'PERFORMANCELENS_DB_PATH',
        os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..', 'performance.db'),
    )


ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.xlsx', '.xls', '.csv'}


# ===========================================================================
# Database
# ===========================================================================

def get_db():
    conn = sqlite3.connect(_db_path())
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = get_db()
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS users (
            id       INTEGER PRIMARY KEY AUTOINCREMENT,
            name     TEXT NOT NULL,
            pin_hash TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS subjects (
            id   INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT UNIQUE NOT NULL
        );
        CREATE TABLE IF NOT EXISTS results (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            term_name   TEXT NOT NULL,
            school_year TEXT NOT NULL DEFAULT '',
            term_label  TEXT NOT NULL DEFAULT '',
            subject     TEXT NOT NULL,
            score       REAL NOT NULL
        );
        CREATE TABLE IF NOT EXISTS marking_scheme (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            grade_label TEXT NOT NULL,
            min_score   REAL NOT NULL,
            max_score   REAL NOT NULL,
            sort_order  INTEGER NOT NULL
        );
    """)

    # Seed default scheme only if table is empty
    existing = conn.execute('SELECT COUNT(*) FROM marking_scheme').fetchone()[0]
    if existing == 0:
        conn.executemany(
            'INSERT INTO marking_scheme (grade_label, min_score, max_score, sort_order) VALUES (?, ?, ?, ?)',
            [
                ('A',  76, 100, 1),
                ('B',  66,  75, 2),
                ('C+', 56,  65, 3),
                ('C-', 46,  55, 4),
                ('D',  40,  45, 5),
                ('F',   0,  39, 6),
            ]
        )
    conn.commit()
    _migrate_results_schema(conn)
    _migrate_study_schema(conn)
    conn.close()


def _migrate_results_schema(conn):
    """Add school_year / term_label to existing DBs; backfill term_label from term_name."""
    cols = {row[1] for row in conn.execute('PRAGMA table_info(results)').fetchall()}
    if 'school_year' not in cols:
        conn.execute("ALTER TABLE results ADD COLUMN school_year TEXT NOT NULL DEFAULT ''")
    if 'term_label' not in cols:
        conn.execute("ALTER TABLE results ADD COLUMN term_label TEXT NOT NULL DEFAULT ''")
    conn.commit()
    conn.execute("""
        UPDATE results
        SET term_label = TRIM(term_name)
        WHERE TRIM(COALESCE(term_label, '')) = ''
          AND TRIM(COALESCE(term_name, '')) != ''
    """)
    conn.commit()


def _migrate_study_schema(conn):
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS study_availability (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL DEFAULT 1,
            day_of_week INTEGER NOT NULL,
            start_minutes INTEGER NOT NULL,
            end_minutes INTEGER NOT NULL
        )
        """
    )
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS study_preferences (
            user_id INTEGER PRIMARY KEY DEFAULT 1,
            weak_subject_multiplier REAL NOT NULL DEFAULT 1.75,
            session_length_min INTEGER NOT NULL DEFAULT 45,
            break_min INTEGER NOT NULL DEFAULT 10,
            max_blocks_per_day INTEGER NOT NULL DEFAULT 8,
            weekend_intensity REAL NOT NULL DEFAULT 1.0,
            reserve_weekly_minutes INTEGER NOT NULL DEFAULT 90,
            min_session_min INTEGER NOT NULL DEFAULT 20,
            updated_at TEXT
        )
        """
    )
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS study_schedules (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL DEFAULT 1,
            title TEXT,
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            plan_json TEXT
        )
        """
    )
    conn.commit()
    n = conn.execute('SELECT COUNT(*) FROM study_preferences').fetchone()[0]
    if n == 0:
        conn.execute(
            """
            INSERT INTO study_preferences (user_id, updated_at) VALUES (1, ?)
            """,
            (datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ'),),
        )
        conn.commit()


def _session_uid():
    return int(session.get('user_id') or 1)


def _make_period_key(school_year, term_label):
    """Stable key for a result period (used by dashboard / analysis JS)."""
    return (school_year or '') + '\x1f' + (term_label or '')


def _collapse_subject_ws(text):
    return re.sub(r'\s+', ' ', (text or '').strip())


def _format_subject_display(name):
    """
    Canonical display form for a new subject: trim, collapse spaces, title case.
    So 'chemistry', 'CHEMISTRY' → 'Chemistry' (when not matching an existing row).
    """
    s = _collapse_subject_ws(name)
    if not s:
        return ''
    return s.title()


def _resolve_subject_name(conn, raw_name):
    """
    Match existing subject case-insensitively (subjects table, then past results).
    Otherwise return title-cased formatted name. Misspellings are not corrected.
    """
    formatted = _format_subject_display(raw_name)
    if not formatted:
        return ''
    key = formatted.casefold()
    for r in conn.execute('SELECT name FROM subjects ORDER BY length(name) DESC').fetchall():
        if r['name'].casefold() == key:
            return r['name']
    for r in conn.execute('SELECT DISTINCT subject FROM results').fetchall():
        if r['subject'].casefold() == key:
            return r['subject']
    return formatted


def _normalize_results_rows(conn, results):
    """
    Resolve subject names, clamp scores to 0–100, drop empty subjects.
    Case-insensitive duplicates collapse; last row wins.
    """
    merged = {}
    for r in results:
        canon = _resolve_subject_name(conn, r.get('subject', ''))
        if not canon:
            continue
        try:
            score = float(r['score'])
        except (TypeError, ValueError):
            continue
        score = max(0.0, min(100.0, score))
        merged[canon.casefold()] = {'subject': canon, 'score': score}
    return list(merged.values())


def _get_subjects():
    """Return all subject names currently stored in the DB."""
    conn = get_db()
    rows = conn.execute('SELECT name FROM subjects').fetchall()
    conn.close()
    return [r['name'] for r in rows]


# ===========================================================================
# File Parsing Helpers
# ===========================================================================

def _find_score(cells):
    """Return the first numeric value between 0–100 found in a list of cell strings."""
    for cell in cells:
        try:
            val = float(str(cell).strip())
            if 0 <= val <= 100:
                return val
        except (ValueError, TypeError):
            continue
    return None


def _detect_term(raw_text):
    """Legacy: any single string that looks like a term line on the report."""
    patterns = [
        r'(Term\s+\d+[\s,]*\d{4})',
        r'(Term\s+\d+)',
        r'(First|Second|Third|Fourth)\s+Term',
        r'(Fall|Spring|Summer|Winter)\s+\d{4}',
        r'(Midterm[s]?|Finals?)',
        r'(\d{4}[-/]\d{4})',
    ]
    for pattern in patterns:
        match = re.search(pattern, raw_text, re.IGNORECASE)
        if match:
            return match.group(0).strip()
    return None


def _detect_school_year(raw_text):
    """Extract academic / school year (e.g. 2024-2025) from document text."""
    if not raw_text:
        return ''
    m = re.search(
        r'\b(20\d{2})\s*[-/]\s*(20\d{2})\b',
        raw_text,
    )
    if m:
        return f'{m.group(1)}-{m.group(2)}'
    m = re.search(
        r'(?:Academic|School)\s+Year\s*:?\s*(20\d{2})\s*[-/]\s*(20\d{2})',
        raw_text,
        re.IGNORECASE,
    )
    if m:
        return f'{m.group(1)}-{m.group(2)}'
    m = re.search(r'(?:Grade|Form)\s+\d+\s*,?\s*(20\d{2})\s*[-/](20\d{2})', raw_text, re.IGNORECASE)
    if m:
        return f'{m.group(1)}-{m.group(2)}'
    return ''


def _detect_term_label(raw_text):
    """Extract term name only (e.g. Term 1, First Term), not the year."""
    if not raw_text:
        return ''
    patterns = [
        r'\b(First|Second|Third|Fourth)\s+Term\b',
        r'\b(Term\s*\d+)\b',
        r'\b(Fall|Spring|Summer|Winter)\s+20\d{2}\b',
        r'\b(Midterm|Midterms|Final|Finals)\b',
    ]
    for pattern in patterns:
        m = re.search(pattern, raw_text, re.IGNORECASE)
        if m:
            return m.group(0).strip()
    m = re.search(r'\bTerm\s+\d+\b', raw_text, re.IGNORECASE)
    if m:
        return m.group(0).strip()
    return ''


def _compose_term_display(school_year, term_label, legacy=None):
    sy = (school_year or '').strip()
    tl = (term_label or '').strip()
    leg = (legacy or '').strip()
    if tl and sy:
        return f'{tl} · {sy}'
    if tl:
        return tl
    if sy:
        return f'Year {sy}'
    return leg or 'Term'


def _build_period_from_text(raw_text):
    """
    Parse school year and term label separately; build display string.
    """
    sy = _detect_school_year(raw_text)
    tl = _detect_term_label(raw_text)
    legacy = _detect_term(raw_text)
    if not tl and legacy:
        # Whole-line match often includes year — keep as label if we have no separate term
        tl = legacy
    td = _compose_term_display(sy, tl, legacy)
    return {
        'school_year': sy,
        'term_label': tl,
        'term_display': td,
    }


def _fuzzy_match_subject(cell_text, subjects):
    """
    Return the matching subject name if cell_text fuzzy-matches any known subject.
    Subjects are checked longest-first so "Further Mathematics" wins over "Mathematics".
    """
    cell_cf = cell_text.casefold()
    for sub in sorted(subjects, key=len, reverse=True):
        s_cf = sub.casefold()
        if s_cf in cell_cf or cell_cf in s_cf:
            return sub
    return None


def _parse_rows(rows, subjects, raw_text=''):
    """
    Match rows against known subjects and extract scores.
    Returns (period_dict, [{"subject": ..., "score": ...}]).
    period_dict has school_year, term_label, term_display.

    If subjects table is empty (e.g. user skipped onboarding upload), falls
    back to blind extraction exactly like ?raw=true — so the input page always
    works regardless of whether subjects were seeded during onboarding.
    """
    # No known subjects to match against — go straight to blind extraction
    if not subjects:
        return _extract_raw(rows, raw_text)

    period = _build_period_from_text(raw_text)
    results = []
    for row in rows:
        str_cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
        for i, cell in enumerate(str_cells):
            subject = _fuzzy_match_subject(cell, subjects)
            if subject:
                score = _find_score(str_cells[i + 1:i + 4])
                if score is not None:
                    results.append({'subject': subject, 'score': score})
                    break

    # Fuzzy matching found nothing — fall back to blind extraction
    # then re-match results against known subjects to keep names consistent
    if not results:
        blind = _extract_from_text_lines(raw_text) if raw_text else []
        if not blind:
            _, blind = _extract_raw(rows, raw_text)
        for item in blind:
            matched = _fuzzy_match_subject(item['subject'], subjects)
            results.append({'subject': matched or item['subject'], 'score': item['score']})

    return period, results


def _extract_from_text_lines(raw_text):
    """
    Regex-based fallback for PDFs where pdfplumber finds no tables.
    Handles the Gambian school report card format (and similar):
      <row_num>  <SUBJECT NAME IN CAPS>  <score>  <grade_letter>
    Also catches plain "Subject Name  85" style lines.
    """
    results = []
    seen = set()

    # Pattern 1: numbered row — "1  CIVIC EDUCATION  83  A"
    numbered = re.compile(
        r'^\d{1,2}[\s.]+([A-Z][A-Z\s&/()\.\-]{2,}?)\s{2,}(\d{1,3})\s+[A-F]',
        re.MULTILINE
    )
    for m in numbered.finditer(raw_text):
        subject = m.group(1).strip()
        score   = float(m.group(2))
        if 0 <= score <= 100 and subject not in seen:
            seen.add(subject)
            results.append({'subject': subject.title(), 'score': score})

    # Pattern 2: plain "Subject Name   85" (Title Case or mixed)
    if not results:
        plain = re.compile(
            r'^([A-Za-z][A-Za-z\s&/()\.\-]{2,}?)\s{2,}(\d{1,3})(?:\s|$)',
            re.MULTILINE
        )
        for m in plain.finditer(raw_text):
            subject = m.group(1).strip()
            score   = float(m.group(2))
            if 0 <= score <= 100 and subject not in seen:
                seen.add(subject)
                results.append({'subject': subject, 'score': score})

    return results


# Metadata cell patterns to ignore during blind extraction
_METADATA_RE = re.compile(
    r'(?i)^(gpa|cgpa|times\s|number\s+of|overall|termly|cumulative|remark|'
    r'result|grade|position|shift|class\b|year\b|student|subject|mark|'
    r'pass|fail|credit|absent|late|detention|suspension|signature)',
)


def _is_metadata(cell):
    return bool(_METADATA_RE.match(cell))


def _extract_raw(all_rows, raw_text=''):
    """
    Extract subject/score pairs without needing known subjects.
    Used during onboarding (?raw=true) when the subjects table is still empty.
    Primary: scan table rows for (text cell, nearby numeric cell) pairs.
    Fallback: regex line scanning for PDFs where table detection fails.
    """
    results = []
    seen = set()

    for row in all_rows:
        str_cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
        if len(str_cells) < 2:
            continue
        for i, cell in enumerate(str_cells):
            # Skip purely numeric cells and known metadata labels
            if re.fullmatch(r'[\d.]+', cell) or _is_metadata(cell):
                continue
            score = _find_score(str_cells[i + 1:i + 4])
            if score is not None and cell not in seen:
                seen.add(cell)
                results.append({'subject': cell.title(), 'score': score})
                break

    # If table extraction found nothing, fall back to text-line parsing
    if not results and raw_text:
        results = _extract_from_text_lines(raw_text)

    return _build_period_from_text(raw_text), results


# ---------------------------------------------------------------------------
# Format-specific parsers
# Each returns (all_rows, raw_text) so the caller can choose _parse_rows
# or _extract_raw depending on context.
# ---------------------------------------------------------------------------

def _read_pdf(file_bytes):
    import pdfplumber
    raw_text, all_rows = '', []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            raw_text += page.extract_text() or ''
            for table in page.extract_tables():
                all_rows.extend(table)
    return all_rows, raw_text


def _read_docx(file_bytes):
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    raw_text = '\n'.join(p.text for p in doc.paragraphs)
    all_rows = [
        [cell.text for cell in row.cells]
        for table in doc.tables
        for row in table.rows
    ]
    return all_rows, raw_text


def _read_excel(file_bytes):
    import pandas as pd
    dfs = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None, header=None)
    all_rows, raw_text = [], ''
    for df in dfs.values():
        for row in df.values.tolist():
            all_rows.append(row)
            raw_text += ' '.join(str(c) for c in row if c) + '\n'
    return all_rows, raw_text


def _read_csv(file_bytes):
    text = file_bytes.decode('utf-8', errors='replace')
    all_rows = list(csv.reader(io.StringIO(text)))
    return all_rows, text


def _read_file(file_bytes, ext):
    """Dispatch to the correct reader based on extension. Returns (all_rows, raw_text)."""
    if ext == '.pdf':
        return _read_pdf(file_bytes)
    elif ext == '.docx':
        return _read_docx(file_bytes)
    elif ext in ('.xlsx', '.xls'):
        return _read_excel(file_bytes)
    elif ext == '.csv':
        return _read_csv(file_bytes)
    else:
        raise ValueError(f'Unsupported file type: {ext}')


# ===========================================================================
# Page Routes
# ===========================================================================

@bp.route('/', endpoint='home')
def home():
    if session.get('user_id') and session.get('pin_unlocked'):
        return redirect(url_for('main.dashboard'))
    if session.get('user_id') and not session.get('pin_unlocked'):
        return redirect(url_for('main.unlock_page'))
    return render_template('home.jinja')


@bp.route('/unlock', endpoint='unlock_page')
def unlock_page():
    if not session.get('user_id'):
        return redirect(url_for('main.home'))
    if session.get('pin_unlocked'):
        return redirect(url_for('main.dashboard'))
    return render_template('unlock.jinja', user_name=(session.get('user_name') or '').strip())


@bp.route('/settings', endpoint='settings')
def settings_page():
    return render_template('settings.jinja', active_page='settings')


@bp.route('/api/logout', methods=['POST'], endpoint='api_logout')
def api_logout():
    session.clear()
    return jsonify({'status': 'success'})


@bp.route('/api/settings/profile', methods=['POST'], endpoint='api_settings_profile')
def api_settings_profile():
    payload = request.get_json(silent=True) or {}
    name = (payload.get('name') or '').strip()
    if not name:
        return jsonify({'message': 'Name is required.'}), 400
    uid = _session_uid()
    conn = get_db()
    cur = conn.execute('UPDATE users SET name = ? WHERE id = ?', (name, uid))
    conn.commit()
    conn.close()
    if cur.rowcount == 0:
        return jsonify({'message': 'User not found.'}), 404
    session['user_name'] = name
    session.modified = True
    return jsonify({'status': 'success', 'name': name})


@bp.route('/api/settings/pin', methods=['POST'], endpoint='api_settings_pin')
def api_settings_pin():
    payload = request.get_json(silent=True) or {}
    cur_pin = (payload.get('current_pin') or '').strip()
    new_pin = (payload.get('new_pin') or '').strip()
    if not cur_pin.isdigit() or len(cur_pin) != 4:
        return jsonify({'message': 'Current PIN must be exactly 4 digits.'}), 400
    if not new_pin.isdigit() or len(new_pin) != 4:
        return jsonify({'message': 'New PIN must be exactly 4 digits.'}), 400
    uid = _session_uid()
    conn = get_db()
    row = conn.execute('SELECT pin_hash FROM users WHERE id = ?', (uid,)).fetchone()
    conn.close()
    if not row or not check_password_hash(row['pin_hash'], cur_pin):
        return jsonify({'message': 'Current PIN is incorrect.'}), 401
    conn = get_db()
    conn.execute(
        'UPDATE users SET pin_hash = ? WHERE id = ?',
        (generate_password_hash(new_pin), uid),
    )
    conn.commit()
    conn.close()
    return jsonify({'status': 'success'})


@bp.route('/api/export/json', methods=['GET'], endpoint='api_export_json')
def api_export_json():
    uid = _session_uid()
    conn = get_db()
    try:
        scheme_rows = conn.execute(
            'SELECT grade_label, min_score, max_score, sort_order FROM marking_scheme ORDER BY sort_order'
        ).fetchall()
        marking_scheme = [dict(r) for r in scheme_rows]
        results_rows = conn.execute(
            'SELECT id, term_name, school_year, term_label, subject, score FROM results ORDER BY id'
        ).fetchall()
        results = [dict(r) for r in results_rows]
        subjects = [r['name'] for r in conn.execute('SELECT name FROM subjects ORDER BY name').fetchall()]
        user_row = conn.execute('SELECT id, name FROM users WHERE id = ?', (uid,)).fetchone()
        user_out = (
            {'id': int(user_row['id']), 'name': user_row['name']}
            if user_row
            else None
        )
        avail = [
            dict(r)
            for r in conn.execute(
                'SELECT day_of_week, start_minutes, end_minutes FROM study_availability WHERE user_id = ? ORDER BY day_of_week, start_minutes',
                (uid,),
            ).fetchall()
        ]
        pref_row = conn.execute('SELECT * FROM study_preferences WHERE user_id = ?', (uid,)).fetchone()
        preferences = None
        if pref_row:
            preferences = {k: pref_row[k] for k in pref_row.keys()}
            preferences.pop('user_id', None)
        schedules = []
        for r in conn.execute(
            'SELECT id, title, created_at, plan_json FROM study_schedules WHERE user_id = ? ORDER BY id',
            (uid,),
        ):
            schedules.append({
                'id': r['id'],
                'title': r['title'],
                'created_at': r['created_at'],
                'plan_json': r['plan_json'],
            })
        payload = {
            'export_format': 'performancelens-v1',
            'exported_at': datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ'),
            'user': user_out,
            'marking_scheme': marking_scheme,
            'subjects': subjects,
            'results': results,
            'study_availability': avail,
            'study_preferences': preferences,
            'study_schedules': schedules,
        }
    finally:
        conn.close()
    body = json.dumps(payload, indent=2)
    return Response(
        body,
        mimetype='application/json',
        headers={
            'Content-Disposition': 'attachment; filename="performancelens-backup.json"',
            'Cache-Control': 'no-store',
        },
    )


@bp.route('/api/settings/reset-data', methods=['POST'], endpoint='api_settings_reset_data')
def api_settings_reset_data():
    payload = request.get_json(silent=True) or {}
    if (payload.get('confirm') or '').strip().upper() != 'RESET':
        return jsonify({'message': 'Send confirm: "RESET" to erase all grades and planner data.'}), 400
    pin = (payload.get('pin') or '').strip()
    if not pin.isdigit() or len(pin) != 4:
        return jsonify({'message': 'PIN must be exactly 4 digits.'}), 400
    uid = _session_uid()
    conn = get_db()
    row = conn.execute('SELECT pin_hash FROM users WHERE id = ?', (uid,)).fetchone()
    if not row or not check_password_hash(row['pin_hash'], pin):
        conn.close()
        return jsonify({'message': 'Incorrect PIN.'}), 401
    try:
        conn.execute('DELETE FROM results')
        conn.execute('DELETE FROM subjects')
        conn.execute('DELETE FROM study_availability WHERE user_id = ?', (uid,))
        conn.execute('DELETE FROM study_preferences WHERE user_id = ?', (uid,))
        conn.execute('DELETE FROM study_schedules WHERE user_id = ?', (uid,))
        conn.commit()
    except Exception as e:
        conn.rollback()
        conn.close()
        return jsonify({'message': str(e)}), 500
    conn.close()
    return jsonify({
        'status': 'success',
        'message': 'All results, subjects, and planner data on this device were removed. Your account and PIN are unchanged.',
    })


@bp.route('/dashboard', endpoint='dashboard')
def dashboard():
    return render_template('dashboard.jinja', active_page='dashboard')


@bp.route('/input', endpoint='input')
def input_page():
    return render_template('input.jinja', active_page='input')


@bp.route('/analysis', endpoint='analysis')
def analysis_page():
    return render_template('analysis.jinja', active_page='analysis')


@bp.route('/recommendations', endpoint='recommendations')
def recommendations_page():
    return render_template('recommendations.jinja', active_page='recommendations')


@bp.route('/study-schedule', endpoint='study_schedule')
def study_schedule_page():
    return redirect(url_for('main.study'))

@bp.route('/study', endpoint='study')
def study_page():
    return render_template('study.jinja')
 

# ===========================================================================
# API — Subjects
# ===========================================================================

@bp.route('/api/user/subjects', endpoint='api_user_subjects')
def api_user_subjects():
    conn = get_db()
    rows = conn.execute('SELECT name FROM subjects ORDER BY name').fetchall()
    conn.close()
    return jsonify({'subjects': [r['name'] for r in rows]})


# ===========================================================================
# API — Upload & Parse
# ===========================================================================

@bp.route('/api/results/upload', methods=['POST'], endpoint='api_results_upload')
def api_results_upload():
    """
    Parse an uploaded file and return a preview of extracted results.
    Does NOT write to the database.

    Query param:
      ?raw=true   Use blind extraction (onboarding — subjects table is empty).
                  Omit or set false for normal use (matches against known subjects).
    """
    if 'file' not in request.files:
        return jsonify({'status': 'error', 'message': 'No file provided.'}), 400

    f = request.files['file']
    ext = os.path.splitext(f.filename)[1].lower()

    if ext not in ALLOWED_EXTENSIONS:
        return jsonify({'status': 'error', 'message': f'Unsupported file type: {ext}'}), 400

    try:
        all_rows, raw_text = _read_file(f.read(), ext)
    except Exception as e:
        return jsonify({'status': 'error', 'message': f'Failed to read file: {str(e)}'}), 500

    raw_mode = request.args.get('raw', 'false').lower() == 'true'

    if raw_mode:
        period, data = _extract_raw(all_rows, raw_text)
    else:
        period, data = _parse_rows(all_rows, _get_subjects(), raw_text)

    if not data:
        return jsonify({
            'status': 'error',
            'message': 'Could not detect a table in this file. Please enter results manually.'
        })

    conn = get_db()
    try:
        merged = {}
        for item in data:
            canon = _resolve_subject_name(conn, item.get('subject', ''))
            if not canon:
                continue
            try:
                sc = float(item['score'])
            except (TypeError, ValueError):
                continue
            sc = max(0.0, min(100.0, sc))
            merged[canon.casefold()] = {'subject': canon, 'score': sc}
        data = list(merged.values())
    finally:
        conn.close()

    if not data:
        return jsonify({
            'status': 'error',
            'message': 'Could not detect a table in this file. Please enter results manually.',
        })

    return jsonify({
        'status': 'success',
        'school_year': period['school_year'],
        'term_label': period['term_label'],
        'term_display': period['term_display'],
        'term_detected': period['term_display'],
        'data': data,
    })


# ===========================================================================
# API — Save Results
# ===========================================================================

@bp.route('/api/results/save', methods=['POST'], endpoint='api_results_save')
def api_results_save():
    """
    Persist confirmed results (manual or upload) to the database.
    Also upserts any new subjects into the subjects table.

    Payload:
      school_year, term_label — identify the period (duplicates rejected unless overwrite=true)
      term_name — display label (optional; defaults from year + term)
      results — array of {subject, score}
      overwrite — if true, replace existing rows for the same school_year + term_label
    Legacy clients may send only term_name; it is stored as term_label with school_year ''.
    """
    payload = request.get_json(silent=True)
    if not payload:
        return jsonify({'status': 'error', 'message': 'Invalid JSON.'}), 400

    results = payload.get('results', [])
    if not results:
        return jsonify({'status': 'error', 'message': 'No results provided.'}), 400

    sy = (payload.get('school_year') or '').strip()
    tl = (payload.get('term_label') or '').strip()
    term_name = (payload.get('term_name') or '').strip()
    overwrite = bool(payload.get('overwrite'))

    if not tl:
        if not term_name:
            return jsonify({
                'status': 'error',
                'message': 'term_label or term_name is required.',
            }), 400
        tl = term_name
    if not term_name:
        term_name = _compose_term_display(sy, tl)

    conn = get_db()
    try:
        results = _normalize_results_rows(conn, results)
        if not results:
            conn.close()
            return jsonify({'status': 'error', 'message': 'No valid subject/score rows to save.'}), 400

        exists = conn.execute(
            'SELECT 1 FROM results WHERE school_year = ? AND term_label = ? LIMIT 1',
            (sy, tl),
        ).fetchone()
        if exists and not overwrite:
            conn.close()
            return jsonify({
                'status': 'conflict',
                'message': 'Results already exist for this school year and term. Send overwrite: true to replace them.',
                'school_year': sy,
                'term_label': tl,
            }), 409

        if exists and overwrite:
            conn.execute(
                'DELETE FROM results WHERE school_year = ? AND term_label = ?',
                (sy, tl),
            )

        conn.executemany(
            'INSERT INTO results (term_name, school_year, term_label, subject, score) VALUES (?, ?, ?, ?, ?)',
            [(term_name, sy, tl, r['subject'], r['score']) for r in results],
        )
        conn.executemany(
            'INSERT OR IGNORE INTO subjects (name) VALUES (?)',
            [(r['subject'],) for r in results],
        )
        conn.commit()
    except Exception as e:
        conn.rollback()
        conn.close()
        return jsonify({'status': 'error', 'message': str(e)}), 500

    conn.close()
    return jsonify({'status': 'success', 'saved': len(results)})


@bp.route('/api/results/periods', methods=['GET'], endpoint='api_results_periods')
def api_results_periods():
    """List saved periods (school year + term) for manage / delete UI."""
    conn = get_db()
    rows = conn.execute(
        """
        SELECT school_year, term_label, MAX(term_name) AS term_name,
               COUNT(*) AS row_count, MIN(id) AS ord
        FROM results
        GROUP BY school_year, term_label
        ORDER BY ord
        """
    ).fetchall()
    conn.close()
    return jsonify({
        'periods': [
            {
                'school_year': r['school_year'],
                'term_label': r['term_label'],
                'term_name': r['term_name'],
                'row_count': r['row_count'],
                'period_key': _make_period_key(r['school_year'], r['term_label']),
            }
            for r in rows
        ],
    })


@bp.route('/api/results/period', methods=['DELETE'], endpoint='api_results_period_delete')
def api_results_period_delete():
    """Delete all results for one school_year + term_label."""
    payload = request.get_json(silent=True) or {}
    sy = (payload.get('school_year') or '').strip()
    tl = (payload.get('term_label') or '').strip()
    if not tl:
        return jsonify({'status': 'error', 'message': 'term_label is required.'}), 400

    conn = get_db()
    try:
        cur = conn.execute(
            'DELETE FROM results WHERE school_year = ? AND term_label = ?',
            (sy, tl),
        )
        conn.commit()
        deleted = cur.rowcount
    except Exception as e:
        conn.rollback()
        conn.close()
        return jsonify({'status': 'error', 'message': str(e)}), 500

    conn.close()
    return jsonify({'status': 'success', 'deleted': deleted})


# ===========================================================================
# API — Onboarding
# ===========================================================================

@bp.route('/api/dev/session', methods=['POST'], endpoint='api_dev_session')
def api_dev_session():
    """
    Development endpoint to create a session automatically.
    This helps bypass the onboarding process for development.
    DISABLED FOR PRODUCTION SECURITY
    """
    import os
    # Only allow in development mode
    if os.environ.get('FLASK_ENV') != 'development':
        return jsonify({'message': 'Development endpoint not available in production'}), 403
    
    # Find or create a test user
    conn = get_db()
    cursor = conn.execute('SELECT id, name FROM users WHERE name = "Test User"')
    user = cursor.fetchone()
    
    if not user:
        # Create test user
        cursor = conn.execute('INSERT INTO users (name, pin_hash) VALUES (?, ?)', 
                           ('Test User', 'scrypt:326872$326872$16$16$salt$hash'))
        user_id = cursor.lastrowid
        user_name = 'Test User'
    else:
        user_id = user['id']
        user_name = user['name']
    
    # Create session
    session['user_id'] = user_id
    session['user_name'] = user_name
    session['pin_unlocked'] = True
    
    conn.close()
    
    return jsonify({
        'status': 'success',
        'message': 'Development session created',
        'user_id': user_id,
        'user_name': user_name
    })


@bp.route('/api/onboarding/complete', methods=['POST'], endpoint='api_onboarding_complete')
def api_onboarding_complete():
    """
    Finalise onboarding. Creates the user, starts a session, and optionally
    saves their first term entry if they uploaded a file during setup.

    Payload (JSON):
      name        string   required
      pin         string   required, exactly 4 digits
      school_year string   optional
      term_label  string   optional (with term_name fallback for legacy)
      term_name   string   optional display — omit if file step was skipped
      results     array    optional — omit if file step was skipped
    """
    payload = request.get_json(silent=True)
    if not payload:
        return jsonify({'message': 'Invalid JSON.'}), 400

    name      = payload.get('name', '').strip()
    pin       = payload.get('pin',  '').strip()
    sy_ob = (payload.get('school_year') or '').strip()
    tl_ob = (payload.get('term_label') or '').strip()
    term_name = (payload.get('term_name') or '').strip()
    results   = payload.get('results', [])

    if not name:
        return jsonify({'message': 'Name is required.'}), 400
    if not pin or not pin.isdigit() or len(pin) != 4:
        return jsonify({'message': 'PIN must be exactly 4 digits.'}), 400

    conn = get_db()

    try:
        cursor = conn.execute(
            'INSERT INTO users (name, pin_hash) VALUES (?, ?)',
            (name, generate_password_hash(pin))
        )
        user_id = cursor.lastrowid
        conn.commit()
    except Exception as e:
        conn.rollback()
        conn.close()
        return jsonify({'message': f'Failed to create user: {str(e)}'}), 500

    session['user_id']   = user_id
    session['user_name'] = name
    session['pin_unlocked'] = True

    # Save first term + seed subjects table — best-effort, never blocks the response
    if results:
        tl_save = tl_ob or term_name
        if not tl_save:
            tl_save = 'First Term'
        if not term_name:
            term_name = _compose_term_display(sy_ob, tl_save)
        try:
            results = _normalize_results_rows(conn, results)
            if results:
                conn.executemany(
                    'INSERT INTO results (term_name, school_year, term_label, subject, score) VALUES (?, ?, ?, ?, ?)',
                    [(term_name, sy_ob, tl_save, r['subject'], r['score']) for r in results],
                )
                conn.executemany(
                    'INSERT OR IGNORE INTO subjects (name) VALUES (?)',
                    [(r['subject'],) for r in results],
                )
                conn.commit()
        except Exception:
            conn.rollback()

    conn.close()
    return jsonify({'status': 'success'})


@bp.route('/api/unlock', methods=['POST'], endpoint='api_unlock')
def api_unlock():
    """Verify PIN for the current session user and mark this session as unlocked."""
    payload = request.get_json(silent=True) or {}
    pin = (payload.get('pin') or '').strip()
    uid = session.get('user_id')
    if not uid:
        return jsonify({'message': 'No account in session.'}), 400
    if not pin or not pin.isdigit() or len(pin) != 4:
        return jsonify({'message': 'PIN must be exactly 4 digits.'}), 400

    conn = get_db()
    row = conn.execute('SELECT pin_hash FROM users WHERE id = ?', (uid,)).fetchone()
    conn.close()
    if not row or not check_password_hash(row['pin_hash'], pin):
        return jsonify({'message': 'Incorrect PIN.'}), 401

    session['pin_unlocked'] = True
    session.modified = True
    return jsonify({'status': 'success'})


# ===========================================================================
# API — Dashboard
# ===========================================================================

def _round1(val):
    return round(val, 1) if val is not None else None


def _marking_scheme_rows(conn):
    return conn.execute(
        'SELECT grade_label, min_score, max_score, sort_order FROM marking_scheme ORDER BY min_score DESC'
    ).fetchall()


def _grade_label_for_score(scheme_rows, score):
    if not scheme_rows or score is None:
        return None
    for row in scheme_rows:
        if score >= row['min_score']:
            return row['grade_label']
    return scheme_rows[-1]['grade_label']


def _next_grade_target(scheme_rows, score):
    """If not already in the top band, return the next higher grade threshold."""
    if not scheme_rows or score is None:
        return None
    for i, row in enumerate(scheme_rows):
        if score >= row['min_score']:
            if i == 0:
                return None
            upper = scheme_rows[i - 1]
            gap = upper['min_score'] - score
            return {
                'grade': upper['grade_label'],
                'min_score': upper['min_score'],
                'gap': _round1(gap) if gap > 0 else 0,
            }
    lowest = scheme_rows[-1]
    if score < lowest['min_score']:
        return {
            'grade': lowest['grade_label'],
            'min_score': lowest['min_score'],
            'gap': _round1(lowest['min_score'] - score),
        }
    return None


def _band_floor_for_score(scheme, score):
    if not scheme or score is None:
        return None, None
    for row in scheme:
        if score >= row['min_score']:
            return float(row['min_score']), row['grade_label']
    return float(scheme[-1]['min_score']), scheme[-1]['grade_label']


def _rec_action_plan(name, avg, grade, next_g, delta_last, weak, volatile, margin_floor):
    steps = []
    if weak or (margin_floor is not None and margin_floor <= 4):
        steps.append({
            'title': 'Stabilize foundations',
            'detail': (
                f'In {name}, spend one session listing every learning objective from the last term you could not explain aloud. '
                f'Prioritize the three that appeared on the weakest assessments.'
            ),
        })
        steps.append({
            'title': 'Error log (ongoing)',
            'detail': (
                f'Keep a running page of mistakes in {name}: question type, what you assumed, correct rule. '
                f'Review the log before each new homework set.'
            ),
        })
    if volatile:
        steps.append({
            'title': 'Even out performance',
            'detail': (
                f'Your marks in {name} swing between terms — use the same pre-test routine every time '
                f'(sleep, 25-minute recap sheet, two warm-up questions) so conditions stay constant.'
            ),
        })
    if next_g and next_g.get('gap', 0) and next_g['gap'] <= 8:
        steps.append({
            'title': 'Target the next grade band',
            'detail': (
                f'You are about {next_g["gap"]} percentage points below {next_g["grade"]}. '
                f'Collect 10 past questions at that difficulty and cycle them twice this week with full worked solutions.'
            ),
        })
    elif next_g and next_g.get('gap'):
        steps.append({
            'title': 'Long arc improvement',
            'detail': (
                f'Closing the gap to {next_g["grade"]} needs steady volume: split {name} into weekly subtopics and '
                f'alternate new material with spaced review of older chapters.'
            ),
        })
    if delta_last is not None and delta_last < -2:
        steps.append({
            'title': 'Recover last term slip',
            'detail': (
                f'{name} dropped {abs(delta_last)} points vs the prior term. Re-do the lowest-scoring assignment or test '
                f'without notes, then compare to the memo — close every gap before moving on.'
            ),
        })
    if not weak and (margin_floor is None or margin_floor > 8):
        steps.append({
            'title': 'Maintain depth',
            'detail': (
                f'{name} is a strength at {avg}%. Add one extension problem or past-paper stretch question per week '
                f'so skills stay sharp for harder terms ahead.'
            ),
        })
    if len(steps) < 3:
        steps.append({
            'title': 'Active recall block',
            'detail': (
                f'For {name}, use closed-book recall: write everything you remember on a topic, then fill gaps from notes. '
                f'Shorter, harder sessions beat long passive reading.'
            ),
        })
    return steps[:6]


def _build_recommendations(conn):
    meta_disclaimer = (
        'Suggestions are educational heuristics based on your saved scores — not tutoring, medical, or professional advice.'
    )
    empty = {
        'has_data': False,
        'summary': 'Add at least one term of results to unlock personalized recommendations.',
        'pulse': None,
        'focus': [],
        'strengths': [],
        'all_ranked': [],
        'shifts': None,
        'shift_leaders': None,
        'consistency': None,
        'volatility_watch': [],
        'at_risk_grades': [],
        'trajectories': [],
        'study_time_suggestion': None,
        'action_plans': [],
        'habits': [],
        'meta_disclaimer': meta_disclaimer,
    }

    results_n = conn.execute('SELECT COUNT(*) FROM results').fetchone()[0]
    if not results_n:
        return empty

    overall = conn.execute('SELECT AVG(score) FROM results').fetchone()[0]
    terms_count = conn.execute(
        'SELECT COUNT(*) FROM (SELECT 1 FROM results GROUP BY school_year, term_label)'
    ).fetchone()[0]
    subjects_count = conn.execute('SELECT COUNT(DISTINCT subject) FROM results').fetchone()[0]
    latest_term_row = conn.execute(
        """
        SELECT school_year, term_label, MAX(term_name) AS term_name
        FROM results
        GROUP BY school_year, term_label
        ORDER BY MIN(id) DESC
        LIMIT 1
        """
    ).fetchone()
    latest_term_name = latest_term_row['term_name'] if latest_term_row else None

    subject_avgs = conn.execute(
        'SELECT subject, AVG(score) AS avg FROM results GROUP BY subject'
    ).fetchall()
    ranked = sorted(subject_avgs, key=lambda r: r['avg'], reverse=True)
    avgs_only = [float(r['avg']) for r in ranked]
    spread = _round1(statistics.stdev(avgs_only)) if len(avgs_only) > 1 else 0.0

    scheme = _marking_scheme_rows(conn)

    def subject_card(row):
        avg = float(row['avg'])
        ng = _next_grade_target(scheme, avg) if scheme else None
        floor, _gl = _band_floor_for_score(scheme, avg) if scheme else (None, None)
        margin = _round1(avg - floor) if floor is not None else None
        return {
            'name': row['subject'],
            'average': _round1(avg),
            'grade': _grade_label_for_score(scheme, avg) if scheme else None,
            'next_grade': ng,
            'margin_above_band_floor': margin,
        }

    focus_n = min(5, len(ranked))
    focus = [subject_card(r) for r in reversed(ranked[-focus_n:])] if ranked else []
    strengths_n = min(5, len(ranked))
    strengths = [subject_card(r) for r in ranked[:strengths_n]] if ranked else []
    all_ranked = [subject_card(r) for r in ranked]
    weak_names = {r['name'] for r in focus}

    term_order = conn.execute(
        """
        SELECT school_year, term_label, MAX(term_name) AS term_name, MIN(id) AS ord
        FROM results
        GROUP BY school_year, term_label
        ORDER BY ord
        """
    ).fetchall()

    subject_series = defaultdict(list)
    for t in term_order:
        rows = conn.execute(
            'SELECT subject, score FROM results WHERE school_year = ? AND term_label = ?',
            (t['school_year'], t['term_label']),
        ).fetchall()
        for r in rows:
            subject_series[r['subject']].append(float(r['score']))

    volatility_watch = []
    trajectories = []
    for sub, series in subject_series.items():
        n = len(series)
        if n >= 2:
            st = _round1(statistics.stdev(series))
            if st >= 6:
                volatility_watch.append({
                    'subject': sub,
                    'stdev': st,
                    'terms': n,
                    'note': 'High term-to-term variance — build consistent revision habits before the next assessment.',
                })
            first, last = series[0], series[-1]
            delta = _round1(last - first)
            if delta >= 3:
                label = 'upward'
            elif delta <= -3:
                label = 'downward'
            else:
                label = 'steady'
            trajectories.append({
                'subject': sub,
                'first_score': _round1(first),
                'latest_score': _round1(last),
                'delta_terms': delta,
                'label': label,
            })
    volatility_watch.sort(key=lambda x: -x['stdev'])
    volatility_watch = volatility_watch[:6]
    trajectories.sort(key=lambda x: -abs(x['delta_terms']))

    at_risk_grades = []
    for r in ranked:
        avg = float(r['avg'])
        floor, gl = _band_floor_for_score(scheme, avg) if scheme else (None, None)
        if floor is None:
            continue
        margin = avg - floor
        if margin <= 3:
            at_risk_grades.append({
                'subject': r['subject'],
                'average': _round1(avg),
                'grade': gl,
                'margin_above_floor': _round1(margin),
                'note': 'Close to the lower edge of this grade band — one weak assessment could pull the average down.',
            })
    at_risk_grades.sort(key=lambda x: x['margin_above_floor'])

    shifts = None
    shift_leaders = None
    deltas = []
    if len(term_order) >= 2:
        prev_t = term_order[-2]
        cur_t = term_order[-1]
        prev_rows = conn.execute(
            'SELECT subject, score FROM results WHERE school_year = ? AND term_label = ?',
            (prev_t['school_year'], prev_t['term_label']),
        ).fetchall()
        cur_rows = conn.execute(
            'SELECT subject, score FROM results WHERE school_year = ? AND term_label = ?',
            (cur_t['school_year'], cur_t['term_label']),
        ).fetchall()
        prev_map = {r['subject']: float(r['score']) for r in prev_rows}
        cur_map = {r['subject']: float(r['score']) for r in cur_rows}
        common = set(prev_map) & set(cur_map)
        deltas = [
            {
                'subject': sub,
                'delta': _round1(cur_map[sub] - prev_map[sub]),
                'latest': _round1(cur_map[sub]),
                'previous': _round1(prev_map[sub]),
            }
            for sub in common
        ]
        deltas.sort(key=lambda x: x['delta'])
        shifts = {
            'previous_term': prev_t['term_name'],
            'latest_term': cur_t['term_name'],
            'biggest_drop': deltas[0] if deltas else None,
            'biggest_gain': deltas[-1] if deltas else None,
            'all_count': len(deltas),
        }
        k = min(5, len(deltas))
        gains_sorted = sorted(deltas, key=lambda x: -x['delta'])
        shift_leaders = {
            'declines': deltas[:k] if deltas else [],
            'gains': gains_sorted[:k] if deltas else [],
        }

    delta_by_subject = {d['subject']: d['delta'] for d in deltas} if deltas else {}

    if spread >= 12:
        consistency_note = 'Very wide spread between subjects — consider balancing time so one weak area does not drag the overall profile.'
    elif spread >= 7:
        consistency_note = 'Moderate spread; a few subjects pull away from the pack.'
    else:
        consistency_note = 'Relatively even profile across subjects.'

    weak_ct = len(focus)
    min_h = 3 + weak_ct * 1.25
    max_h = 5 + weak_ct * 2
    study_time_suggestion = {
        'min_hours_week': _round1(min_h),
        'max_hours_week': _round1(max_h),
        'rationale': (
            f'Heuristic based on {weak_ct} priority subject(s) and {terms_count} term period(s) of data. '
            f'Adjust for school workload and wellbeing.'
        ),
    }

    action_plans = []
    for card in focus + strengths[:2]:
        name = card['name']
        volatile = any(v['subject'] == name for v in volatility_watch)
        steps = _rec_action_plan(
            name,
            card['average'],
            card['grade'],
            card.get('next_grade'),
            delta_by_subject.get(name),
            name in weak_names,
            volatile,
            card.get('margin_above_band_floor'),
        )
        action_plans.append({
            'subject': name,
            'urgency': 'high' if name in weak_names else 'maintain',
            'average': card['average'],
            'grade': card['grade'],
            'steps': steps,
        })

    habits = [
        {
            'title': 'Same sleep window before test weeks',
            'body': 'Keeping wake/sleep times within 60 minutes day-to-day stabilizes recall more than last-minute cramming.',
        },
        {
            'title': 'Teach-back Friday',
            'body': 'Each Friday, explain one hard topic aloud for 5 minutes as if to a classmate — gaps show where notes are thin.',
        },
        {
            'title': 'Phone-away deep blocks',
            'body': 'Match your study schedule deep blocks with phone-in-another-room to protect the minutes you already committed.',
        },
    ]

    summary = (
        f'Overall {_round1(overall)}% across {terms_count} term period(s) and {subjects_count} subject(s). '
        f'Latest data: {latest_term_name or "—"}. {consistency_note}'
    )

    pulse = {
        'overall_average': _round1(overall),
        'terms_count': terms_count,
        'subjects_count': subjects_count,
        'results_rows': results_n,
        'latest_term': latest_term_name,
        'spread_across_subjects': spread,
    }

    return {
        'has_data': True,
        'summary': summary,
        'pulse': pulse,
        'focus': focus,
        'strengths': strengths,
        'all_ranked': all_ranked,
        'shifts': shifts,
        'shift_leaders': shift_leaders,
        'consistency': {
            'stdev_across_subject_averages': spread,
            'note': consistency_note,
        },
        'volatility_watch': volatility_watch,
        'at_risk_grades': at_risk_grades,
        'trajectories': trajectories[:10],
        'study_time_suggestion': study_time_suggestion,
        'action_plans': action_plans,
        'habits': habits,
        'meta_disclaimer': meta_disclaimer,
    }


@bp.route('/api/recommendations', endpoint='api_recommendations')
def api_recommendations():
    conn = get_db()
    try:
        payload = _build_recommendations(conn)
    finally:
        conn.close()
    return jsonify(payload)


@bp.route('/api/dashboard/summary', endpoint='api_dashboard_summary')
def api_dashboard_summary():
    conn = get_db()

    overall = conn.execute('SELECT AVG(score) FROM results').fetchone()[0]
    subjects_count = conn.execute('SELECT COUNT(*) FROM subjects').fetchone()[0]
    terms_count = conn.execute(
        'SELECT COUNT(*) FROM (SELECT 1 FROM results GROUP BY school_year, term_label)'
    ).fetchone()[0]
    results_count = conn.execute('SELECT COUNT(*) FROM results').fetchone()[0]

    subject_avgs = conn.execute(
        'SELECT subject, AVG(score) as avg FROM results GROUP BY subject'
    ).fetchall()

    conn.close()

    best = worst = None
    if subject_avgs:
        ranked = sorted(subject_avgs, key=lambda r: r['avg'], reverse=True)
        best  = {'name': ranked[0]['subject'],  'average': _round1(ranked[0]['avg'])}
        worst = {'name': ranked[-1]['subject'], 'average': _round1(ranked[-1]['avg'])}

    return jsonify({
        'overall_average': _round1(overall),
        'subjects_count':  subjects_count,
        'terms_count':     terms_count,
        'results_count':   results_count,
        'best_subject':    best,
        'worst_subject':   worst,
    })


@bp.route('/api/dashboard/trend', endpoint='api_dashboard_trend')
def api_dashboard_trend():
    conn = get_db()
    rows = conn.execute(
        """
        SELECT school_year, term_label, MAX(term_name) AS term_name,
               AVG(score) AS avg, MIN(id) AS ord
        FROM results
        GROUP BY school_year, term_label
        ORDER BY ord
        """
    ).fetchall()
    conn.close()

    return jsonify({
        'trend': [
            {
                'term_name': r['term_name'],
                'school_year': r['school_year'],
                'term_label': r['term_label'],
                'period_key': _make_period_key(r['school_year'], r['term_label']),
                'average': _round1(r['avg']),
            }
            for r in rows
        ],
    })


@bp.route('/api/dashboard/terms', endpoint='api_dashboard_terms')
def api_dashboard_terms():
    """
    All recorded terms in chronological order (by first insert), each with
    per-subject scores and term-level aggregates for dashboard term cards.
    """
    conn = get_db()
    term_order = conn.execute(
        """
        SELECT school_year, term_label, MAX(term_name) AS term_name, MIN(id) AS ord
        FROM results
        GROUP BY school_year, term_label
        ORDER BY ord
        """
    ).fetchall()

    terms = []
    for row in term_order:
        sy, tl = row['school_year'], row['term_label']
        display = row['term_name']
        rows = conn.execute(
            """
            SELECT subject, score FROM results
            WHERE school_year = ? AND term_label = ?
            ORDER BY score DESC
            """,
            (sy, tl),
        ).fetchall()
        results = [{'subject': r['subject'], 'score': r['score']} for r in rows]
        avg = conn.execute(
            'SELECT AVG(score) FROM results WHERE school_year = ? AND term_label = ?',
            (sy, tl),
        ).fetchone()[0]
        hi = conn.execute(
            'SELECT MAX(score) FROM results WHERE school_year = ? AND term_label = ?',
            (sy, tl),
        ).fetchone()[0]
        lo = conn.execute(
            'SELECT MIN(score) FROM results WHERE school_year = ? AND term_label = ?',
            (sy, tl),
        ).fetchone()[0]
        terms.append({
            'term_name': display,
            'school_year': sy,
            'term_label': tl,
            'period_key': _make_period_key(sy, tl),
            'average': _round1(avg),
            'high': _round1(hi),
            'low': _round1(lo),
            'subject_count': len(results),
            'results': results,
        })

    conn.close()
    return jsonify({'terms': terms})


@bp.route('/api/dashboard/subjects', endpoint='api_dashboard_subjects')
def api_dashboard_subjects():
    conn = get_db()

    subject_rows = conn.execute(
        'SELECT subject, AVG(score) as avg FROM results GROUP BY subject ORDER BY avg DESC'
    ).fetchall()

    all_scores = conn.execute(
        'SELECT subject, school_year, term_label, score FROM results'
    ).fetchall()

    term_order = conn.execute(
        """
        SELECT school_year, term_label, MAX(term_name) AS term_name, MIN(id) AS ord
        FROM results
        GROUP BY school_year, term_label
        ORDER BY ord
        """
    ).fetchall()
    term_sequence = [
        _make_period_key(r['school_year'], r['term_label']) for r in term_order
    ]
    term_display = {
        _make_period_key(r['school_year'], r['term_label']): r['term_name']
        for r in term_order
    }

    scores_map = defaultdict(dict)
    for row in all_scores:
        pk = _make_period_key(row['school_year'], row['term_label'])
        scores_map[row['subject']][pk] = row['score']

    subjects = []
    for row in subject_rows:
        sub = row['subject']
        subjects.append({
            'name':    sub,
            'average': _round1(row['avg']),
            'scores_by_term': [
                {
                    'period_key': pk,
                    'term_name': term_display[pk],
                    'score': scores_map[sub][pk],
                }
                for pk in term_sequence
                if pk in scores_map[sub]
            ]
        })

    latest_term = None
    if term_order:
        last = term_order[-1]
        sy, tl = last['school_year'], last['term_label']
        pk = _make_period_key(sy, tl)
        lt_rows = conn.execute(
            """
            SELECT subject, score FROM results
            WHERE school_year = ? AND term_label = ?
            ORDER BY score DESC
            """,
            (sy, tl),
        ).fetchall()
        latest_term = {
            'term_name': last['term_name'],
            'school_year': sy,
            'term_label': tl,
            'period_key': pk,
            'results': [{'subject': r['subject'], 'score': r['score']} for r in lt_rows],
        }

    conn.close()
    return jsonify({'subjects': subjects, 'latest_term': latest_term})


# ===========================================================================
# API — Marking Scheme
# ===========================================================================

@bp.route('/api/settings/marking-scheme', endpoint='api_marking_scheme_get')
def api_marking_scheme_get():
    conn = get_db()
    rows = conn.execute(
        'SELECT grade_label, min_score, max_score FROM marking_scheme ORDER BY sort_order'
    ).fetchall()
    conn.close()
    return jsonify({
        'scheme': [
            {'grade': r['grade_label'], 'min': r['min_score'], 'max': r['max_score']}
            for r in rows
        ]
    })


@bp.route('/api/settings/marking-scheme', methods=['PUT'], endpoint='api_marking_scheme_put')
def api_marking_scheme_put():
    """
    Replace the entire marking scheme. Payload:
    {
        "scheme": [
            { "grade": "A",  "min": 76, "max": 100 },
            { "grade": "B",  "min": 66, "max": 75  },
            ...
        ]
    }
    Grades must be ordered highest-first (descending by min_score).
    """
    payload = request.get_json(silent=True)
    if not payload or 'scheme' not in payload:
        return jsonify({'message': 'Invalid payload.'}), 400

    scheme = payload['scheme']
    if not isinstance(scheme, list) or len(scheme) == 0:
        return jsonify({'message': 'scheme must be a non-empty array.'}), 400

    for entry in scheme:
        if not all(k in entry for k in ('grade', 'min', 'max')):
            return jsonify({'message': 'Each entry needs grade, min, and max.'}), 400

    conn = get_db()
    try:
        conn.execute('DELETE FROM marking_scheme')
        conn.executemany(
            'INSERT INTO marking_scheme (grade_label, min_score, max_score, sort_order) VALUES (?, ?, ?, ?)',
            [(e['grade'], e['min'], e['max'], i + 1) for i, e in enumerate(scheme)]
        )
        conn.commit()
    except Exception as e:
        conn.rollback()
        conn.close()
        return jsonify({'message': str(e)}), 500

    conn.close()
    return jsonify({'status': 'success'})


# ===========================================================================
# API — Study schedule
# ===========================================================================

def _prefs_from_row(row):
    base = sched.default_preferences()
    if not row:
        return base
    base.update({
        'weak_subject_multiplier': float(row['weak_subject_multiplier']),
        'session_length_min': int(row['session_length_min']),
        'break_min': int(row['break_min']),
        'max_blocks_per_day': int(row['max_blocks_per_day']),
        'weekend_intensity': float(row['weekend_intensity']),
        'reserve_weekly_minutes': int(row['reserve_weekly_minutes']),
        'min_session_min': int(row['min_session_min']),
    })
    return base


_STUDY_PREF_KEYS = (
    'weak_subject_multiplier',
    'session_length_min',
    'break_min',
    'max_blocks_per_day',
    'weekend_intensity',
    'reserve_weekly_minutes',
    'min_session_min',
)


def _merge_study_prefs_from_payload(base: dict, payload: dict) -> dict:
    """Overlay JSON body onto prefs with strict int/float coercion (avoids 500s from bad types)."""
    out = dict(base)
    for key in _STUDY_PREF_KEYS:
        if key not in payload:
            continue
        raw = payload[key]
        try:
            if key in ('weak_subject_multiplier', 'weekend_intensity'):
                out[key] = float(raw)
            else:
                out[key] = int(raw)
        except (TypeError, ValueError) as e:
            raise ValueError(f'Invalid value for {key!r} — use a number.') from e
    return out


@bp.route('/api/study/availability', methods=['GET', 'PUT'], endpoint='api_study_availability')
def api_study_availability():
    uid = _session_uid()
    conn = get_db()
    try:
        if request.method == 'GET':
            rows = conn.execute(
                """
                SELECT day_of_week, start_minutes, end_minutes
                FROM study_availability WHERE user_id = ? ORDER BY day_of_week, start_minutes
                """,
                (uid,),
            ).fetchall()
            return jsonify({'availability': sched.db_rows_to_availability_json(rows)})

        payload = request.get_json(silent=True) or {}
        raw = payload.get('availability') or payload
        by_day = sched.normalize_availability_payload(raw)
        conn.execute('DELETE FROM study_availability WHERE user_id = ?', (uid,))
        for row in sched.availability_to_db_rows(uid, by_day):
            conn.execute(
                """
                INSERT INTO study_availability (user_id, day_of_week, start_minutes, end_minutes)
                VALUES (?, ?, ?, ?)
                """,
                row,
            )
        conn.commit()
        return jsonify({'status': 'success'})
    finally:
        conn.close()


@bp.route('/api/study/preferences', methods=['GET', 'PUT'], endpoint='api_study_preferences')
def api_study_preferences():
    uid = _session_uid()
    conn = get_db()
    try:
        if request.method == 'GET':
            row = conn.execute(
                'SELECT * FROM study_preferences WHERE user_id = ?', (uid,)
            ).fetchone()
            if not row:
                conn.execute(
                    'INSERT INTO study_preferences (user_id, updated_at) VALUES (?, ?)',
                    (uid, datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')),
                )
                conn.commit()
                row = conn.execute(
                    'SELECT * FROM study_preferences WHERE user_id = ?', (uid,)
                ).fetchone()
            return jsonify({'preferences': _prefs_from_row(row)})

        payload = request.get_json(silent=True) or {}
        p = _prefs_from_row(
            conn.execute('SELECT * FROM study_preferences WHERE user_id = ?', (uid,)).fetchone()
        )
        try:
            p = _merge_study_prefs_from_payload(p, payload)
        except ValueError as e:
            return jsonify({'status': 'error', 'message': str(e)}), 400
        conn.execute(
            """
            INSERT INTO study_preferences (
                user_id, weak_subject_multiplier, session_length_min, break_min,
                max_blocks_per_day, weekend_intensity, reserve_weekly_minutes,
                min_session_min, updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(user_id) DO UPDATE SET
                weak_subject_multiplier = excluded.weak_subject_multiplier,
                session_length_min = excluded.session_length_min,
                break_min = excluded.break_min,
                max_blocks_per_day = excluded.max_blocks_per_day,
                weekend_intensity = excluded.weekend_intensity,
                reserve_weekly_minutes = excluded.reserve_weekly_minutes,
                min_session_min = excluded.min_session_min,
                updated_at = excluded.updated_at
            """,
            (
                uid,
                float(p['weak_subject_multiplier']),
                int(p['session_length_min']),
                int(p['break_min']),
                int(p['max_blocks_per_day']),
                float(p['weekend_intensity']),
                int(p['reserve_weekly_minutes']),
                int(p['min_session_min']),
                datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ'),
            ),
        )
        conn.commit()
        return jsonify({'status': 'success', 'preferences': p})
    finally:
        conn.close()


@bp.route('/api/study/generate', methods=['POST'], endpoint='api_study_generate')
def api_study_generate():
    uid = _session_uid()
    payload = request.get_json(silent=True) or {}
    title = (payload.get('title') or '').strip()
    if not title:
        return jsonify({
            'status': 'error',
            'message': 'Enter a name for your plan (this appears as the PDF title).',
        }), 400

    conn = get_db()
    try:
        if isinstance(payload.get('availability'), dict):
            by_day = sched.normalize_availability_payload(payload['availability'])
        else:
            rows = conn.execute(
                'SELECT day_of_week, start_minutes, end_minutes FROM study_availability WHERE user_id = ?',
                (uid,),
            ).fetchall()
            by_day = sched.normalize_availability_payload(sched.db_rows_to_availability_json(rows))

        if sched.total_free_minutes_week(by_day) < 45:
            return jsonify({
                'status': 'error',
                'message': 'Add at least ~45 minutes of weekly availability before generating.',
            }), 400

        sub_rows = conn.execute(
            'SELECT subject, AVG(score) AS avg FROM results GROUP BY subject'
        ).fetchall()
        if not sub_rows:
            return jsonify({
                'status': 'error',
                'message': 'No subject results in the database — add scores under Manage results first.',
            }), 400

        if isinstance(payload.get('availability'), dict):
            conn.execute('DELETE FROM study_availability WHERE user_id = ?', (uid,))
            for row in sched.availability_to_db_rows(uid, by_day):
                conn.execute(
                    """
                    INSERT INTO study_availability (user_id, day_of_week, start_minutes, end_minutes)
                    VALUES (?, ?, ?, ?)
                    """,
                    row,
                )

        ranked_asc = sorted(sub_rows, key=lambda r: float(r['avg']))
        weak_names = {r['subject'] for r in ranked_asc[: min(5, len(ranked_asc))]}
        subjects_ranked = [
            {'name': r['subject'], 'average': _round1(float(r['avg']))}
            for r in sorted(sub_rows, key=lambda r: float(r['avg']), reverse=True)
        ]

        pref_row = conn.execute(
            'SELECT * FROM study_preferences WHERE user_id = ?', (uid,)
        ).fetchone()
        prefs = _prefs_from_row(pref_row)
        try:
            prefs = _merge_study_prefs_from_payload(prefs, payload)
        except ValueError as e:
            return jsonify({'status': 'error', 'message': str(e)}), 400

        plan = sched.build_plan(subjects_ranked, weak_names, by_day, prefs)
        student = (session.get('user_name') or '').strip()
        now = datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')
        plan_json = sched.plan_to_json_str(plan)
        cur = conn.execute(
            """
            INSERT INTO study_schedules (user_id, title, created_at, plan_json)
            VALUES (?, ?, ?, ?)
            """,
            (uid, title, now, plan_json),
        )
        sid = cur.lastrowid
        conn.commit()
        return jsonify({
            'status': 'success',
            'id': sid,
            'plan': plan,
            'preview_url': url_for('main.api_study_schedule_one', schedule_id=sid),
        })
    except Exception as e:
        conn.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({
            'status': 'error',
            'message': str(e) or 'Server error while generating the schedule.',
            'error_type': type(e).__name__,
        }), 500
    finally:
        conn.close()


@bp.route('/api/study/schedules', methods=['GET'], endpoint='api_study_schedules_list')
def api_study_schedules_list():
    uid = _session_uid()
    conn = get_db()
    rows = conn.execute(
        """
        SELECT id, title, created_at,
               LENGTH(plan_json) AS plan_chars
        FROM study_schedules WHERE user_id = ? ORDER BY id DESC
        """,
        (uid,),
    ).fetchall()
    conn.close()
    return jsonify({
        'schedules': [
            {
                'id': r['id'],
                'title': r['title'],
                'created_at': r['created_at'],
                'preview_url': url_for('main.api_study_schedule_one', schedule_id=r['id']),
            }
            for r in rows
        ],
    })


@bp.route('/api/study/schedules/<int:schedule_id>', methods=['GET', 'DELETE'], endpoint='api_study_schedule_one')
def api_study_schedule_one(schedule_id):
    uid = _session_uid()
    conn = get_db()
    try:
        row = conn.execute(
            'SELECT * FROM study_schedules WHERE id = ? AND user_id = ?',
            (schedule_id, uid),
        ).fetchone()
        if not row:
            conn.close()
            return jsonify({'message': 'Schedule not found.'}), 404
        if request.method == 'DELETE':
            conn.execute('DELETE FROM study_schedules WHERE id = ? AND user_id = ?', (schedule_id, uid))
            conn.commit()
            conn.close()
            return jsonify({'status': 'success', 'message': 'Schedule deleted successfully.'})
        plan = json.loads(row['plan_json'])
        conn.close()
        return jsonify({'id': row['id'], 'title': row['title'], 'created_at': row['created_at'], 'plan': plan})
    except Exception as e:
        conn.close()
        return jsonify({'message': f'Failed to load schedule: {str(e)}'}), 500



